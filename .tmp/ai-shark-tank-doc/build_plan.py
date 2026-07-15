from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "outputs" / "ai-shark-tank-implementation-plan"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUT_DIR / "AI_Shark_Tank_Implementation_Plan.docx"

NAVY = "14213D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GOLD = "B7791F"
WHITE = "FFFFFF"
BORDER = "D0D5DD"

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles

def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
]:
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.keep_with_next = True

title_style = styles["Title"]
title_style.font.name = "Calibri"
title_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
title_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
title_style.font.size = Pt(30)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor.from_string(NAVY)
title_style.paragraph_format.space_before = Pt(0)
title_style.paragraph_format.space_after = Pt(8)

subtitle_style = styles["Subtitle"]
subtitle_style.font.name = "Calibri"
subtitle_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
subtitle_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
subtitle_style.font.size = Pt(15)
subtitle_style.font.color.rgb = RGBColor.from_string(DARK_BLUE)
subtitle_style.paragraph_format.space_before = Pt(0)
subtitle_style.paragraph_format.space_after = Pt(6)

if "Callout" not in styles:
    callout_style = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
else:
    callout_style = styles["Callout"]
callout_style.font.name = "Calibri"
callout_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
callout_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
callout_style.font.size = Pt(11)
callout_style.font.color.rgb = RGBColor.from_string(INK)
callout_style.paragraph_format.left_indent = Inches(0.18)
callout_style.paragraph_format.right_indent = Inches(0.18)
callout_style.paragraph_format.space_before = Pt(8)
callout_style.paragraph_format.space_after = Pt(10)
callout_style.paragraph_format.line_spacing = 1.25

def shade_paragraph(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), BORDER)
        borders.append(el)
    pPr.append(borders)

def add_callout(label, text):
    p = doc.add_paragraph(style="Callout")
    shade_paragraph(p, CALLOUT)
    r = p.add_run(label + " ")
    set_font(r, size=11, color=NAVY, bold=True)
    r = p.add_run(text)
    set_font(r, size=11, color=INK)
    return p

def create_numbering_definition(kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    pPr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    pPr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    pPr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    lvl.append(pPr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id

BULLET_NUM_ID = create_numbering_definition("bullet")
DECIMAL_NUM_ID = create_numbering_definition("decimal")

def apply_numbering(p, num_id):
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    apply_numbering(p, BULLET_NUM_ID)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, size=11)
    return p

def add_number(text):
    p = doc.add_paragraph(style="List Number")
    apply_numbering(p, DECIMAL_NUM_ID)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, size=11)
    return p

def configure_list_styles():
    for style_name in ["List Bullet", "List Number"]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.188)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25
    st = styles["List Bullet 2"]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(11)
    st.paragraph_format.left_indent = Inches(0.75)
    st.paragraph_format.first_line_indent = Inches(-0.188)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.25

configure_list_styles()

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[i] / 1440)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[i]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def add_table(headers, rows, widths, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(str(text))
        set_font(r, size=10, color=NAVY, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if alignments:
                p.alignment = alignments[i]
            r = p.add_run(str(value))
            set_font(r, size=9.5, color="000000")
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table

def add_footer_with_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("AI Shark Tank | Implementation Plan | ")
    set_font(r, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)

def add_header(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("AI SHARK TANK")
    set_font(r, size=8.5, color=MUTED, bold=True)

add_header(section)
add_footer_with_page_number(section)

# COVER - editorial_cover named override
for _ in range(5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
r = p.add_run("PRODUCT AND ENGINEERING ROADMAP")
set_font(r, size=10.5, color=GOLD, bold=True)

p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("AI Shark Tank")

p = doc.add_paragraph(style="Subtitle")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Implementation Plan for AI Evaluation, Credits, Billing, Privacy, and Quality")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(72)
r = p.add_run("A staged plan for moving from the current MVP to a reliable paid beta")
set_font(r, size=10.5, color=MUTED, italic=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(3)
r = p.add_run("Prepared for AI Shark Tank")
set_font(r, size=12, color=NAVY, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("July 13, 2026 | Version 1.0")
set_font(r, size=9.5, color=MUTED)

doc.add_page_break()

# EXECUTIVE SUMMARY
doc.add_heading("Executive summary", level=1)
p = doc.add_paragraph()
p.add_run("Objective. ").bold = True
p.add_run("Add a reliable two-tier AI pitch evaluation system, automated video and deck processing, credit-based monetization, privacy controls, and systematic quality evaluation to the existing authenticated MVP.")

add_callout(
    "Core recommendation:",
    "Do not train a model for the first release. Launch with DeepSeek Flash for the free report and a multi-evaluator DeepSeek Pro workflow for premium reports. Improve quality through prompts, retrieval, examples, validation, and evaluations before considering fine-tuning.",
)

doc.add_heading("Current foundation", level=2)
p = doc.add_paragraph("The current application already provides the essential product shell:")
for item in [
    "Supabase authentication and private storage buckets for videos and decks.",
    "A structured submission form and saved report dashboard.",
    "A typed investor report format with scores, questions, risks, milestones, and valuation framing.",
    "A single-call report generator with demo fallback behavior.",
]:
    add_bullet(item)

doc.add_heading("Principal gaps", level=2)
for item in [
    "Automated transcription and deck-text extraction.",
    "A background worker for processing beyond the current synchronous request limit.",
    "Separate basic and premium evaluation pipelines.",
    "Credits, purchases, subscriptions, and idempotent payment webhooks.",
    "Data-retention controls, consent records, abuse prevention, and operational monitoring.",
    "A repeatable evaluation dataset for prompt and model changes.",
]:
    add_bullet(item)

doc.add_heading("Recommended beta offer", level=1)
pricing_rows = [
    ["Free", "$0", "1 lifetime basic pitch", "DeepSeek V4 Flash"],
    ["Pitch Pack", "$7.99 once", "5 premium pitches", "DeepSeek V4 Pro"],
    ["Builder", "$9.99/month", "15 premium pitches", "DeepSeek V4 Pro"],
    ["Additional Pack", "$4.99", "5 premium pitches", "DeepSeek V4 Pro"],
]
add_table(
    ["Product", "Price", "Included usage", "Model"],
    pricing_rows,
    [1900, 1600, 3300, 2560],
    [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

doc.add_heading("Credit rules", level=2)
for item in [
    "Each completed report consumes one applicable credit.",
    "A technical failure returns the reserved credit automatically.",
    "A revised pitch submission consumes another credit.",
    "Purchased credits expire after 12 months.",
    "Subscription credits roll over to a maximum balance of 30.",
    "Free users receive one lifetime basic report, not one free report every month.",
    "Do not offer an unlimited plan during beta.",
]:
    add_bullet(item)

add_callout("Pricing note:", "Treat this pricing as a beta hypothesis. Conversion, resubmission behavior, and retention should determine the final package and limits.")

# PRODUCT WORKFLOW
doc.add_heading("Target product workflow", level=1)
workflow_steps = [
    "The founder signs in, selects Basic or Premium, and uploads a five-minute pitch video and a deck of no more than 20 pages.",
    "The server creates a queued submission and atomically reserves the required credit.",
    "A background worker extracts audio, transcribes the pitch, and extracts deck text.",
    "The worker invokes the appropriate AI pipeline and validates the resulting structured JSON.",
    "The application records model usage, prompt version, token cost, and processing time.",
    "The report is saved, the reserved credit is consumed, and temporary video data is deleted according to policy.",
    "If processing fails, the job records the error and returns the credit.",
]
for step in workflow_steps:
    add_number(step)

doc.add_heading("Processing-state model", level=2)
add_table(
    ["State", "Meaning", "User experience"],
    [
        ["Queued", "Submission accepted and credit reserved", "Show position or waiting status"],
        ["Processing", "Files are being extracted and evaluated", "Show stage-based progress"],
        ["Complete", "Validated report saved", "Open report and consume credit"],
        ["Failed", "Processing could not complete", "Show retry path and return credit"],
    ],
    [1500, 3600, 4260],
    [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

# PHASES
doc.add_heading("Implementation phases", level=1)

phases = [
    ("Phase 1 - Product rules and entitlements", [
        "Encode the free, pack, and subscription offers as server-side product definitions.",
        "Define what counts as a credit, when it is reserved, and when it is returned.",
        "Centralize video duration, deck page, file size, and account-level usage limits.",
    ], "A written entitlement contract and test cases for every credit transition."),
    ("Phase 2 - Asynchronous processing", [
        "Change submission creation to save a queued record rather than generating a report in the request.",
        "Add a background job with retry, timeout, and failure handling.",
        "Expose a status endpoint and poll or subscribe from the report screen.",
    ], "A submission survives browser closure and completes independently of the web request."),
    ("Phase 3 - Automated ingestion", [
        "Extract audio from supported video formats.",
        "Transcribe no more than five minutes of audio.",
        "Extract PDF text locally and add a PPTX extraction or conversion path.",
        "Reject encrypted, corrupt, or over-limit files with clear errors.",
    ], "The model receives a transcript and normalized deck text without requiring manual paste-in."),
    ("Phase 4 - AI provider layer", [
        "Introduce a provider interface so model configuration is isolated from product logic.",
        "Configure DeepSeek V4 Flash for basic reports and DeepSeek V4 Pro for premium reports.",
        "Record model name, prompt version, tokens, estimated cost, duration, and retry count.",
    ], "Models can be swapped through configuration without rewriting report or billing code."),
    ("Phase 5 - Basic and premium reports", [
        "Basic: one Flash call with a concise score, strengths, risks, questions, and next steps.",
        "Premium: five parallel Pro evaluator calls covering economics, product, go-to-market, brand, and execution risk.",
        "Premium synthesis: a sixth Pro call reconciles the five reviews into one evidence-based report.",
    ], "A premium report is visibly more detailed and useful than the free report."),
    ("Phase 6 - Reliability and validation", [
        "Validate every response with Zod and enforce scores from 0 to 100.",
        "Retry malformed JSON once and never silently replace a failed premium report with demo content.",
        "Require evidence grounding and prevent invented traction, revenue, or financial claims.",
    ], "At least 95% of test reports complete without manual repair."),
    ("Phase 7 - Credits and billing", [
        "Create products, subscriptions, credit ledger, payment events, model runs, and report version records.",
        "Use Stripe Checkout and Customer Portal for purchases and account management.",
        "Grant credits only from verified, idempotently processed webhooks.",
        "Reserve and consume credits atomically; return them automatically after failure.",
    ], "No duplicate webhook or concurrent request can create or spend the same credit twice."),
    ("Phase 8 - Product interface", [
        "Add pricing, billing, available-credit, and processing-status screens.",
        "Show the exact credit cost before submission and confirmation before a paid resubmission.",
        "Add Basic or Premium labels, report export, pitch history, and version comparison.",
    ], "A founder always understands the cost, status, and value level of a submission."),
    ("Phase 9 - Privacy and abuse controls", [
        "Collect processing consent and keep model-improvement consent separate and opt-in.",
        "Do not log raw pitch content; retain only what the product needs.",
        "Delete raw videos after processing or after a short published retention window.",
        "Verify email before granting the free report and rate-limit submissions.",
    ], "Users can understand, control, and delete their data; free access is difficult to abuse."),
    ("Phase 10 - Evaluation and model improvement", [
        "Create a private set of 30 to 50 pitches across stages and industries.",
        "Measure grounding, question quality, actionability, consistency, JSON reliability, latency, and cost.",
        "Version prompts and run the test set before any model or prompt release.",
        "Consider fine-tuning only after collecting a large, explicitly permitted, human-corrected dataset.",
    ], "Model changes are approved by evidence rather than subjective spot checks."),
]

for heading, tasks, outcome in phases:
    doc.add_heading(heading, level=2)
    for task in tasks:
        add_bullet(task)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Completion signal: ")
    set_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r = p.add_run(outcome)
    set_font(r, size=10.5, color=INK, italic=True)

# DATA MODEL
doc.add_heading("Data model additions", level=1)
add_table(
    ["Table", "Purpose"],
    [
        ["products", "Server-side definitions for prices, included credits, model tier, and limits."],
        ["subscriptions", "User subscription state, provider identifiers, renewal date, and cancellation status."],
        ["credit_ledger", "Immutable grants, reservations, consumption, returns, expirations, and refunds."],
        ["payment_events", "Idempotent record of processed Stripe webhook events."],
        ["model_runs", "Provider, model, prompt version, tokens, estimated cost, duration, and result status."],
        ["report_versions", "Links repeated pitches and supports version-to-version comparison."],
        ["user_consents", "Processing, retention, and optional model-improvement permissions."],
    ],
    [2700, 6660],
    [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

doc.add_heading("Credit ledger principles", level=2)
for item in [
    "Never calculate entitlement only from browser state.",
    "Never mutate a single balance without an auditable ledger event.",
    "Use an idempotency key for submission reservations and payment events.",
    "Make reservation and queued-submission creation one database transaction.",
    "Associate refunds and reversals with the original grant.",
]:
    add_bullet(item)

# MODEL PIPELINE
doc.add_heading("AI evaluation design", level=1)
doc.add_heading("Basic report", level=2)
p = doc.add_paragraph("One DeepSeek Flash request receives the profile, transcript, deck text, scoring rubric, and strict JSON schema. It returns:")
for item in ["Overall score", "Three strengths", "Three risks", "Five investor questions", "Three next steps", "A concise executive summary"]:
    add_bullet(item)

doc.add_heading("Premium report", level=2)
evaluator_rows = [
    ["1", "Profitability and unit economics", "Revenue logic, margins, use of funds, and financial assumptions"],
    ["2", "Product and defensibility", "Problem-solution fit, differentiation, technical moat, and execution"],
    ["3", "Founder and go-to-market", "Founder-market fit, acquisition plan, traction, and sales motion"],
    ["4", "Brand and customer appeal", "Positioning, narrative, community, trust, and customer value"],
    ["5", "Scale and investment risk", "Operational risks, market expansion, capital needs, and milestones"],
]
add_table(
    ["Evaluator", "Lens", "Primary focus"],
    evaluator_rows,
    [1100, 3000, 5260],
    [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)
p = doc.add_paragraph("A sixth Pro request synthesizes the five independent reviews into one coherent report. Evaluator names should be original and should not impersonate or imply endorsement by real investors.")

doc.add_heading("Required report metadata", level=2)
for item in [
    "Tier: basic or premium",
    "AI provider and exact model",
    "Prompt version",
    "Input and output tokens",
    "Estimated variable cost",
    "Processing time and retry count",
    "Evidence-grounding status",
]:
    add_bullet(item)

# CODE MAP
doc.add_heading("Recommended code changes", level=1)
code_rows = [
    ["app/api/submissions/route.ts", "Create queued submissions, reserve credits, and enqueue processing."],
    ["lib/report-generator.ts", "Refactor into basic and premium orchestration using a provider interface."],
    ["lib/ai/deepseek.ts", "DeepSeek-compatible client, model configuration, usage parsing, and retries."],
    ["lib/ai/schemas.ts", "Zod schemas for evaluator outputs and final reports."],
    ["lib/processing/*", "Audio extraction, transcription, PDF/PPTX text extraction, and cleanup."],
    ["lib/billing/*", "Credit ledger operations, Stripe checkout, webhook verification, and reconciliation."],
    ["supabase/schema.sql", "Add credit, subscription, payment, model-run, version, and consent tables."],
    ["components/PitchSubmissionForm.tsx", "Tier selection, credit display, consent, limits, and progress."],
    ["components/ReportView.tsx", "Tier badge, metadata disclosure, comparison, practice, and export actions."],
]
add_table(
    ["Location", "Change"],
    code_rows,
    [3200, 6160],
    [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

# SPRINTS
doc.add_heading("Suggested sprint sequence", level=1)
sprint_rows = [
    ["Sprint 1", "Core AI pipeline", "DeepSeek provider; basic report; five Pro evaluators; synthesis; validation; cost tracking"],
    ["Sprint 2", "Automated processing", "Background jobs; transcription; deck extraction; progress; cleanup and failure handling"],
    ["Sprint 3", "Monetization", "Credit ledger; free credit; Stripe checkout; webhooks; packs; subscription; refunds"],
    ["Sprint 4", "Quality and launch", "Version comparison; pricing UI; privacy controls; rate limits; evaluation dataset; monitoring"],
]
add_table(
    ["Sprint", "Theme", "Primary deliverables"],
    sprint_rows,
    [1250, 2400, 5710],
    [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

doc.add_heading("Recommended implementation order", level=2)
for item in [
    "Prove that the premium report is reliably and visibly better than the basic report.",
    "Automate transcription and deck extraction behind a background job.",
    "Add the credit ledger and failure-safe reservation flow.",
    "Connect payments only after premium quality and job reliability are acceptable.",
    "Run a small paid beta before expanding plan complexity.",
]:
    add_number(item)

# LAUNCH CHECKLIST
doc.add_heading("Paid beta launch checklist", level=1)
check_rows = [
    ["[ ]", "At least 95% of test reports complete without manual intervention."],
    ["[ ]", "Invalid JSON is automatically rejected, repaired, or retried."],
    ["[ ]", "No submission can reserve or consume more than one credit accidentally."],
    ["[ ]", "Failed jobs return reserved credits automatically."],
    ["[ ]", "Every report records model, prompt version, tokens, cost, and duration."],
    ["[ ]", "Raw videos are deleted according to the published retention policy."],
    ["[ ]", "Stripe webhook processing is verified and idempotent."],
    ["[ ]", "The evaluation set meets the agreed grounding and usefulness threshold."],
    ["[ ]", "Ten external founders complete the full workflow successfully."],
    ["[ ]", "Users can delete their pitches, reports, and associated uploaded files."],
]
add_table(
    ["Status", "Launch requirement"],
    check_rows,
    [1100, 8260],
    [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
)

doc.add_heading("Metrics for the first paid beta", level=2)
for item in [
    "Free-to-paid conversion",
    "Reports completed per paid user",
    "Resubmission rate",
    "Median report cost and processing time",
    "Failure and credit-return rate",
    "Pitch Pack versus subscription selection",
    "Subscription retention after the founder's immediate pitch event",
    "Founder usefulness rating and referral rate",
]:
    add_bullet(item)

add_callout("Decision gate:", "Do not introduce more plans or fine-tuning until the beta reveals how often founders revise pitches, which premium features drive payment, and whether ongoing coaching creates genuine subscription retention.")

# REFERENCES
doc.add_heading("Reference links", level=1)
refs = [
    ("DeepSeek models and pricing", "https://api-docs.deepseek.com/quick_start/pricing/"),
    ("Stripe standard pricing", "https://stripe.com/pricing"),
    ("Current AI Shark Tank report generator", "lib/report-generator.ts"),
    ("Current submission API", "app/api/submissions/route.ts"),
    ("Current database schema", "supabase/schema.sql"),
]
for label, url in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label + ": ")
    set_font(r, size=10, color=NAVY, bold=True)
    r = p.add_run(url)
    set_font(r, size=10, color=BLUE)

# Core document properties and save
doc.core_properties.title = "AI Shark Tank Implementation Plan"
doc.core_properties.subject = "AI evaluation, credits, billing, privacy, and quality roadmap"
doc.core_properties.author = "AI Shark Tank"
doc.core_properties.keywords = "AI Shark Tank, DeepSeek, implementation plan, subscription, credits, privacy"

doc.save(OUTPUT)
print(OUTPUT)
